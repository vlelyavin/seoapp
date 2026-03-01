"""HTML report generator."""

import copy
import logging
from datetime import datetime
from html import unescape
from pathlib import Path
from urllib.request import Request, urlopen
from urllib.parse import urlparse

logger = logging.getLogger(__name__)

import re
import unicodedata
from jinja2 import Environment, FileSystemLoader
from markupsafe import Markup, escape

from .config import settings
from .i18n import get_translator, load_translations, _
from .models import AnalyzerResult, AuditResult, SeverityLevel
from .utils import extract_domain

# Singleton instance for ReportGenerator
_report_generator_instance = None

# DOCX export fonts should stay cross-platform to avoid missing-font warnings
DOCX_DEFAULT_FONT = "Arial"
DOCX_MONO_FONT = "Courier New"


def translate_analyzer_content(result: AnalyzerResult, lang: str, translator) -> AnalyzerResult:
    """
    Translate analyzer result content to target language.

    This function handles translation at render time, allowing the analyzer
    code to remain in English (source language) while supporting multiple
    output languages (uk, ru).
    """
    import re

    if lang == 'en':
        return result  # English is the source language

    # Create a deep copy to avoid modifying the original
    translated = copy.deepcopy(result)
    name = result.name

    # Translate theory
    theory_key = f"analyzer_content.{name}.theory"
    translated_theory = translator.get(theory_key, "")
    if translated_theory:
        translated.theory = translated_theory

    # Translate display_name
    display_name_key = f"analyzer_content.{name}.display_name"
    translated_display = translator.get(display_name_key, "")
    if translated_display:
        translated.display_name = translated_display

    # Translate description
    desc_key = f"analyzer_content.{name}.description"
    translated_desc = translator.get(desc_key, "")
    if translated_desc:
        translated.description = translated_desc

    # For some analyzers we rebuild problem summaries from translated issues
    # later in this function to avoid mixed-language fragments.
    rebuild_problem_summary_from_issues = False

    # Translate summary - handle special cases for different analyzers
    if result.summary:
        if name == "cms":
            # Extract CMS name from English summary
            cms_match = re.search(r'using (.+)$', result.summary)
            if cms_match:
                cms_name = cms_match.group(1)
                summary_key = f"analyzer_content.{name}.summary.cms_detected"
                translated_summary = translator.get(summary_key, "")
                if translated_summary and "{cms}" in translated_summary:
                    translated.summary = translated_summary.format(cms=cms_name)
            elif "could not be identified" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.cms_unknown"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary

        elif name == "meta_tags":
            # Handle meta_tags summary with dynamic numbers
            if "Missing meta tags:" in result.summary:
                match = re.search(r'Missing meta tags: (\d+) Title, (\d+) Description', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.missing"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(
                            missing_titles=match.group(1),
                            missing_descriptions=match.group(2)
                        )
                        # Handle duplicates part if present
                        dup_match = re.search(r'Duplicates: (\d+) Title, (\d+) Description', result.summary)
                        if dup_match:
                            dup_key = f"analyzer_content.{name}.summary.duplicates"
                            dup_trans = translator.get(dup_key, "")
                            if dup_trans:
                                translated.summary += ". " + dup_trans.format(
                                    duplicate_titles=dup_match.group(1),
                                    duplicate_descriptions=dup_match.group(2)
                                )
            elif "Duplicates:" in result.summary:
                dup_match = re.search(r'Duplicates: (\d+) Title, (\d+) Description', result.summary)
                if dup_match:
                    dup_key = f"analyzer_content.{name}.summary.duplicates"
                    translated_summary = translator.get(dup_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(
                            duplicate_titles=dup_match.group(1),
                            duplicate_descriptions=dup_match.group(2)
                        )
            elif "have correct meta tags" in result.summary:
                match = re.search(r'All (\d+) pages', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.all_ok"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(total_pages=match.group(1))

        elif name == "headings":
            if "Issues found:" in result.summary:
                match = re.search(r'Issues found: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.problems_found"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(problems=match.group(1))
            elif "have a correct H1" in result.summary:
                match = re.search(r'All (\d+) pages', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.all_ok"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(total_pages=match.group(1))

        elif name == "page_404":
            summary_map = {
                "configured correctly": "ok",
                "created or fixed": "missing",
                "needs improvement": "needs_improvement",
                "Could not check": "check_failed"
            }
            for eng_text, key in summary_map.items():
                if eng_text in result.summary:
                    summary_key = f"analyzer_content.{name}.summary.{key}"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary
                    break

        elif name == "speed":
            # Extract mobile and desktop scores
            match = re.search(r'Mobile: (\d+)/100, Desktop: (\d+)/100', result.summary)
            if match:
                mobile, desktop = match.group(1), match.group(2)
                if "within normal range" in result.summary:
                    key = "ok"
                elif "Optimization needed" in result.summary:
                    key = "needs_optimization"
                elif "Critical" in result.summary:
                    key = "critical"
                else:
                    key = None
                if key:
                    summary_key = f"analyzer_content.{name}.summary.{key}"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(mobile=mobile, desktop=desktop)
            elif "Could not retrieve" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.failed"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary

        elif name == "images":
            if "are optimized" in result.summary:
                match = re.search(r'All (\d+) images', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.all_ok"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(count=match.group(1))
            elif "Found" in result.summary and "images" in result.summary:
                match = re.search(r'Found (\d+) images\. Issues: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.problems"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(count=match.group(1), problems=match.group(2))

        elif name == "content":
            if "have sufficient content" in result.summary:
                match = re.search(r'All (\d+) pages.*Average word count: (\d+)', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.all_ok"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(total_pages=match.group(1), avg_words=match.group(2))
            elif "Content issues:" in result.summary:
                match = re.search(r'Content issues: (.+)\. Average word count: (\d+)', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.problems"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(problems=match.group(1), avg_words=match.group(2))

        elif name == "links":
            if "No issues found" in result.summary:
                match = re.search(r'Checked (\d+) internal and (\d+) external', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.no_broken"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(internal=match.group(1), external=match.group(2))
            elif "Broken links found:" in result.summary:
                rebuild_problem_summary_from_issues = True

        elif name == "favicon":
            summary_map = {
                "configured correctly": "ok",
                "is missing": "missing",
                "can be improved": "needs_improvement"
            }
            for eng_text, key in summary_map.items():
                if eng_text in result.summary:
                    summary_key = f"analyzer_content.{name}.summary.{key}"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary
                    break

        elif name == "external_links":
            match = re.search(r'Found (\d+) external links to (\d+) domains', result.summary)
            if match:
                summary_key = f"analyzer_content.{name}.summary.ok"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary.format(count=match.group(1), domains=match.group(2))
            else:
                match = re.search(r'Found (\d+) external links\. Warnings: (\d+), info: (\d+)', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.with_warnings"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(count=match.group(1), warnings=match.group(2), info=match.group(3))

        elif name == "robots":
            if "in order" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.ok"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary
            elif "Indexation issues:" in result.summary:
                rebuild_problem_summary_from_issues = True

        elif name == "structure":
            if "is optimal" in result.summary:
                match = re.search(r'Maximum depth: (\d+) levels', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.ok"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(depth=match.group(1))
            else:
                match = re.search(r'Maximum depth: (\d+)\. Issues: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.problems"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(depth=match.group(1), problems=match.group(2))

        elif name == "content_sections":
            if "Detected:" in result.summary:
                match = re.search(r'Detected: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.detected"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(sections=match.group(1))
            elif "No informational sections" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.not_detected"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary

        elif name == "schema":
            if "Found" in result.summary and "Schema.org types" in result.summary:
                match = re.search(r'Found (\d+) Schema\.org types across (\d+) pages', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.found"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(types=match.group(1), pages=match.group(2))
            elif "is missing" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.missing"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary

        elif name == "social_tags":
            if "OG tags:" in result.summary:
                match = re.search(r'OG tags: (\d+)/(\d+).*Twitter Cards: (\d+)/(\d+)', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.stats"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(
                            og=match.group(1), total=match.group(2),
                            twitter=match.group(3), total2=match.group(4))
            elif "No pages" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.no_pages"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary

        elif name == "security":
            if "in order" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.ok"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary
            elif "Issues found:" in result.summary:
                rebuild_problem_summary_from_issues = True

        elif name == "mobile":
            if "have a viewport" in result.summary:
                match = re.search(r'All (\d+) pages', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.ok"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(count=match.group(1))
            elif "Issues:" in result.summary:
                match = re.search(r'Issues: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.problems"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(problems=match.group(1))

        elif name == "url_quality":
            if "All URLs are well-structured" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.ok"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary
            elif "Issues found:" in result.summary:
                match = re.search(r'Issues found: (.+)$', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.problems"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(problems=match.group(1))

        elif name == "hreflang":
            if "are missing" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.missing"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary
            elif "Found" in result.summary and "language versions" in result.summary:
                match = re.search(r'Found (\d+) language versions across (\d+) pages', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.found"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(langs=match.group(1), pages=match.group(2))

        elif name == "duplicates":
            if "Found" in result.summary and "duplicate groups" in result.summary:
                match = re.search(r'Found (\d+) duplicate groups', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.found"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(count=match.group(1))
            elif "No duplicates" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.ok"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary

        elif name == "redirects":
            if "Found" in result.summary and "redirect chains" in result.summary:
                match = re.search(r'Found (\d+) redirect chains', result.summary)
                if match:
                    summary_key = f"analyzer_content.{name}.summary.found"
                    translated_summary = translator.get(summary_key, "")
                    if translated_summary:
                        translated.summary = translated_summary.format(count=match.group(1))
            elif "No redirect issues" in result.summary:
                summary_key = f"analyzer_content.{name}.summary.ok"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary

        # Post-process: replace remaining English words in summary
        # (injected via {problems}, {broken}, {sections} placeholders)
        if translated.summary:
            _summary_word_map = {
                'multiple H1': {'uk': 'декілька H1', 'ru': 'несколько H1'},
                'duplicate H1': {'uk': 'дублів H1', 'ru': 'дублей H1'},
                'hierarchy violations': {'uk': 'порушень ієрархії', 'ru': 'нарушений иерархии'},
                'no H1': {'uk': 'без H1', 'ru': 'без H1'},
                'too large': {'uk': 'завеликі', 'ru': 'слишком большие'},
                'outdated format': {'uk': 'застарілий формат', 'ru': 'устаревший формат'},
                'empty': {'uk': 'порожніх', 'ru': 'пустых'},
                'with thin content': {'uk': 'з малим контентом', 'ru': 'с малым контентом'},
                'internal': {'uk': 'внутрішніх', 'ru': 'внутренних'},
                'external': {'uk': 'зовнішніх', 'ru': 'внешних'},
                'deep pages': {'uk': 'глибоких сторінок', 'ru': 'глубоких страниц'},
                'orphan': {'uk': 'сирітських', 'ru': 'сиротских'},
                'errors': {'uk': 'помилок', 'ru': 'ошибок'},
                'warnings': {'uk': 'попереджень', 'ru': 'предупреждений'},
                'no viewport': {'uk': 'без viewport', 'ru': 'без viewport'},
                'Flash content': {'uk': 'Flash-контент', 'ru': 'Flash-контент'},
                'incorrect viewport': {'uk': 'некоректний viewport', 'ru': 'некорректный viewport'},
                'long URLs': {'uk': 'довгих URL', 'ru': 'длинных URL'},
                'uppercase letters': {'uk': 'великі літери', 'ru': 'заглавные буквы'},
                'special characters': {'uk': 'спецсимволи', 'ru': 'спецсимволы'},
                'underscores': {'uk': 'підкреслення', 'ru': 'подчёркивания'},
                'double slashes': {'uk': 'подвійні слеші', 'ru': 'двойные слэши'},
                'parameters': {'uk': 'параметри', 'ru': 'параметры'},
            }
            for eng, translations in _summary_word_map.items():
                if eng in translated.summary and lang in translations:
                    translated.summary = translated.summary.replace(eng, translations[lang])

    # Translate issues
    for issue in translated.issues:
        # Try to translate message by category
        msg_key = f"analyzer_content.{name}.issues.{issue.category}"
        translated_msg = translator.get(msg_key, "")

        if translated_msg:
            try:
                # Handle CMS-specific translations with dynamic CMS name
                if name == "cms" and issue.category == "cms_detected":
                    cms_match = re.search(r'using (.+)$', issue.message)
                    if cms_match and "{cms}" in translated_msg:
                        issue.message = translated_msg.format(cms=cms_match.group(1))
                elif name == "cms" and issue.category == "multiple_cms":
                    cms_match = re.search(r'detected: (.+)$', issue.message)
                    if cms_match and "{cms_list}" in translated_msg:
                        issue.message = translated_msg.format(cms_list=cms_match.group(1))
                # Try to format with count if available (only {count}, no other placeholders)
                elif issue.count is not None and "{count}" in translated_msg and "{" not in translated_msg.replace("{count}", ""):
                    issue.message = translated_msg.format(count=issue.count)
                elif "{" not in translated_msg:
                    # No placeholders, use as-is
                    issue.message = translated_msg
                else:
                    # General fallback: extract dynamic values from English message
                    format_kwargs = {}
                    numbers = re.findall(r'\d+', issue.message)

                    if issue.count is not None:
                        format_kwargs['count'] = issue.count
                    elif '{count}' in translated_msg:
                        # Try to extract ratio like (85/100) first
                        ratio_match = re.search(r'\((\d+/\d+)\)', issue.message)
                        if ratio_match:
                            format_kwargs['count'] = ratio_match.group(1)
                        elif numbers:
                            format_kwargs['count'] = numbers[0]

                    # Extract domain for external_links many_links_same_domain
                    if '{domain}' in translated_msg:
                        domain_match = re.search(r'to (.+?):', issue.message)
                        if domain_match:
                            format_kwargs['domain'] = domain_match.group(1)
                        # Also try to get count from "domain: N"
                        count_match = re.search(r':\s*(\d+)', issue.message)
                        if count_match:
                            format_kwargs['count'] = count_match.group(1)

                    if format_kwargs:
                        issue.message = translated_msg.format(**format_kwargs)
            except (KeyError, ValueError, IndexError):
                # If formatting fails, keep original message
                pass

        # Translate details - handle CMS special case
        details_key = f"analyzer_content.{name}.details.{issue.category}"
        translated_details = translator.get(details_key, "")
        if translated_details:
            if name == "cms" and issue.category == "cms_detected":
                # Extract evidence from original details
                evidence_match = re.search(r'indicators: (.+)$', issue.details or "")
                if evidence_match and "{evidence}" in translated_details:
                    evidence = evidence_match.group(1)
                    issue.details = translated_details.format(evidence=evidence)
            elif "{" not in translated_details:
                issue.details = translated_details

        # Translate recommendation
        rec_key = f"analyzer_content.{name}.recommendations.{issue.category}"
        translated_rec = translator.get(rec_key, "")
        if translated_rec and "{" not in translated_rec:
            issue.recommendation = translated_rec

    # Post-process: replace remaining English words in issue messages
    # (speed FCP/LCP/CLS metrics, content_sections missing features, etc.)
    if translated.issues:
        _issue_word_map = {
            'slow': {'uk': 'повільний', 'ru': 'медленный'},
            'high': {'uk': 'високий', 'ru': 'высокий'},
            'target': {'uk': 'ціль', 'ru': 'цель'},
            'missing elements:': {'uk': 'відсутні елементи:', 'ru': 'отсутствуют элементы:'},
            'publication dates': {'uk': 'дати публікації', 'ru': 'даты публикации'},
            'categories': {'uk': 'категорії', 'ru': 'категории'},
        }
        for issue in translated.issues:
            if issue.message:
                for eng, translations in _issue_word_map.items():
                    if eng in issue.message and lang in translations:
                        issue.message = issue.message.replace(eng, translations[lang])

    # Rebuild selected problem summaries from already translated issue messages
    # to avoid mixed-language output (e.g., "Найдено ... broken links").
    if rebuild_problem_summary_from_issues:
        translated_problem_messages = [
            issue.message for issue in translated.issues
            if issue.message and issue.severity != SeverityLevel.SUCCESS
        ][:3]

        if translated_problem_messages:
            if name == "links":
                summary_key = "analyzer_content.links.summary.broken_found"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary.format(
                        broken=", ".join(translated_problem_messages)
                    )
            elif name in {"robots", "security"}:
                summary_key = f"analyzer_content.{name}.summary.problems"
                translated_summary = translator.get(summary_key, "")
                if translated_summary:
                    translated.summary = translated_summary.format(
                        problems=", ".join(translated_problem_messages)
                    )

    # Translate tables
    # Build reverse maps: English value → target translation (keyed by snake_case keys)
    en_translations = load_translations("en")
    en_tt = en_translations.get("table_translations", {})
    target_tt = translator.translations.get("table_translations", {})

    def build_reverse_map(en_section: dict, target_section: dict) -> dict:
        result = {}
        for key in en_section:
            if key in target_section:
                result[en_section[key]] = target_section[key]
        return result

    table_titles = build_reverse_map(en_tt.get("titles", {}), target_tt.get("titles", {}))
    table_headers = build_reverse_map(en_tt.get("headers", {}), target_tt.get("headers", {}))
    table_values = build_reverse_map(en_tt.get("values", {}), target_tt.get("values", {}))
    table_patterns = build_reverse_map(en_tt.get("patterns", {}), target_tt.get("patterns", {}))

    analyzer_en_table_title = (
        en_translations.get("analyzer_content", {})
        .get(name, {})
        .get("issues", {})
        .get("table_title")
    )
    analyzer_localized_table_title = translator.get(
        f"analyzer_content.{name}.issues.table_title",
        "",
    )

    for table in translated.tables:
        # Translate table title
        if table.get("title"):
            if table["title"] in table_titles:
                table["title"] = table_titles[table["title"]]
            elif (
                analyzer_en_table_title
                and analyzer_localized_table_title
                and table["title"] == analyzer_en_table_title
            ):
                table["title"] = analyzer_localized_table_title

        # Translate table headers
        if table.get("headers"):
            table["headers"] = [
                table_headers.get(h, h) for h in table["headers"]
            ]

        # Translate row values
        if table.get("rows"):
            new_rows = []
            for row in table["rows"]:
                new_row = {}
                for key, value in row.items():
                    # Translate key if it's in table_headers
                    new_key = table_headers.get(key, key)
                    if isinstance(value, str):
                        # 1. Exact match in values map
                        if value in table_values:
                            new_row[new_key] = table_values[value]
                        # 2. Exact match in headers map
                        elif value in table_headers:
                            new_row[new_key] = table_headers[value]
                        # 3. Pattern-based replacement for dynamic strings
                        else:
                            translated_value = value
                            for pattern, replacement in table_patterns.items():
                                if pattern in translated_value:
                                    translated_value = translated_value.replace(pattern, replacement)
                            new_row[new_key] = translated_value
                    else:
                        new_row[new_key] = value
                new_rows.append(new_row)
            table["rows"] = new_rows

    return translated


class ReportGenerator:
    """Generates autonomous HTML reports from audit results."""

    def __init__(self):
        templates_path = Path(__file__).parent / "templates"
        self.env = Environment(
            loader=FileSystemLoader(str(templates_path)),
            autoescape=True,
        )

        # Add custom filters
        self.env.filters['status_icon'] = self.status_icon
        self.env.filters['severity_class'] = self.severity_class
        self.env.filters['format_number'] = self.format_number
        self.env.filters['format_cell'] = self.format_cell

    @staticmethod
    def status_icon(severity: SeverityLevel) -> str:
        """Convert severity to inline SVG icon."""
        icons = {
            SeverityLevel.SUCCESS: '<svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#10b981" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M22 11.08V12a10 10 0 1 1-5.93-9.14"/><polyline points="22 4 12 14.01 9 11.01"/></svg>',
            SeverityLevel.WARNING: '<svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#f59e0b" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="m21.73 18-8-14a2 2 0 0 0-3.48 0l-8 14A2 2 0 0 0 4 21h16a2 2 0 0 0 1.73-3Z"/><line x1="12" y1="9" x2="12" y2="13"/><line x1="12" y1="17" x2="12.01" y2="17"/></svg>',
            SeverityLevel.ERROR: '<svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#ef4444" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><line x1="15" y1="9" x2="9" y2="15"/><line x1="9" y1="9" x2="15" y2="15"/></svg>',
            SeverityLevel.INFO: '<svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="#3b82f6" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><line x1="12" y1="16" x2="12" y2="12"/><line x1="12" y1="8" x2="12.01" y2="8"/></svg>',
        }
        return icons.get(severity, '')

    @staticmethod
    def severity_class(severity: SeverityLevel) -> str:
        """Convert severity to CSS class."""
        classes = {
            SeverityLevel.SUCCESS: 'success',
            SeverityLevel.WARNING: 'warning',
            SeverityLevel.ERROR: 'error',
            SeverityLevel.INFO: 'info',
        }
        return classes.get(severity, 'info')

    @staticmethod
    def format_number(value: int) -> str:
        """Format number with thousands separator."""
        return f"{value:,}".replace(",", " ")

    @staticmethod
    def format_cell(value) -> Markup:
        """Format table cell: replace ✓/✗/⚠️ with SVG icons, make URLs clickable."""
        text = str(value) if value is not None else ""

        # HTML-escape the text first to prevent XSS
        text = str(escape(text))

        # SVG icons (14×14) matching the report's existing icon style
        icon_check = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#10b981" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><polyline points="9 12 12 15 16 10"/></svg>'
        icon_cross = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#ef4444" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><line x1="15" y1="9" x2="9" y2="15"/><line x1="9" y1="9" x2="15" y2="15"/></svg>'
        icon_warning = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#f59e0b" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="m21.73 18-8-14a2 2 0 0 0-3.48 0l-8 14A2 2 0 0 0 4 21h16a2 2 0 0 0 1.73-3Z"/><line x1="12" y1="9" x2="12" y2="13"/><line x1="12" y1="17" x2="12.01" y2="17"/></svg>'

        # Replace Unicode symbols and emojis with SVG icons
        text = text.replace("\u2713", icon_check)     # ✓
        text = text.replace("\u2714", icon_check)     # ✔
        text = text.replace("\u2717", icon_cross)     # ✗
        text = text.replace("\u2718", icon_cross)     # ✘
        text = text.replace("\u2716", icon_cross)     # ✖
        text = text.replace("\u26a0\ufe0f", icon_warning)  # ⚠️ (with variation selector)
        text = text.replace("\u26a0", icon_warning)   # ⚠ (without variation selector)

        # Make URLs clickable
        text = re.sub(
            r'(https?://[^\s<>&"\']+)',
            r'<a href="\1" target="_blank" rel="noopener noreferrer">\1</a>',
            text,
        )

        return Markup(text)

    @staticmethod
    def _build_report_heading(t, brand: dict | None = None) -> str:
        """Return localized top title with optional company name."""
        company_name = ((brand or {}).get("company_name") or "").strip()
        if company_name:
            titled = t("report.website_seo_audit_by", company=company_name)
            if titled != "report.website_seo_audit_by":
                return titled
            return f"Website SEO audit by {company_name}"

        plain = t("report.website_seo_audit")
        if plain != "report.website_seo_audit":
            return plain

        legacy = t("report.express_title")
        if legacy != "report.express_title":
            return legacy
        return "Website SEO audit"

    async def generate(self, audit: AuditResult, brand: dict | None = None) -> str:
        """Generate HTML report and return file path."""
        template = self.env.get_template("report.html")

        # Get translator for the audit language
        lang = getattr(audit, 'language', 'en') or 'en'
        t = get_translator(lang)

        # Prepare sections for navigation with translated names
        sections = []
        section_order = [
            "cms", "speed", "meta_tags", "headings", "page_404",
            "images", "content", "links", "favicon", "external_links",
            "robots", "structure", "content_sections",
            "schema", "social_tags", "security", "mobile",
            "url_quality", "hreflang", "duplicates", "redirects",
        ]

        for name in section_order:
            if name in audit.results:
                result = audit.results[name]

                # Translate analyzer content if not English
                if lang != 'en':
                    result = translate_analyzer_content(result, lang, t)

                # Get translated title, fallback to display_name from result
                title = t(f"analyzers.{name}.name")
                if title == f"analyzers.{name}.name":
                    title = result.display_name  # Fallback to analyzer's display_name

                sections.append({
                    "id": name,
                    "title": title,
                    "icon": "",
                    "severity": result.severity,
                    "result": result,
                })

        # Extract domain
        domain = extract_domain(audit.url)

        # Compute category overview data (sorted by severity)
        severity_order = {SeverityLevel.ERROR: 0, SeverityLevel.WARNING: 1, SeverityLevel.INFO: 2, SeverityLevel.SUCCESS: 3}
        badge_text_map = {
            SeverityLevel.SUCCESS: t("report.badge_ok"),
            SeverityLevel.WARNING: t("report.badge_warning"),
            SeverityLevel.ERROR: t("report.badge_error"),
            SeverityLevel.INFO: t("report.badge_info"),
        }
        category_overview = []
        for section in sorted(sections, key=lambda s: severity_order.get(s["severity"], 4)):
            result = section["result"]
            criticals = sum(1 for iss in result.issues if iss.severity == SeverityLevel.ERROR)
            warns = sum(1 for iss in result.issues if iss.severity == SeverityLevel.WARNING)
            category_overview.append({
                "title": section["title"],
                "severity": section["severity"],
                "badge_text": badge_text_map.get(section["severity"], "—"),
                "criticals": criticals,
                "warns": warns,
            })

        # Prepare translations for template
        translations = {
            "report_title": t("report.title"),
            "express_title": t("report.express_title"),
            "overview": t("report.overview"),
            "pages_crawled": t("report.pages_crawled"),
            "passed_checks": t("report.passed_checks"),
            "warnings": t("report.warnings"),
            "critical_issues": t("report.critical_issues"),
            "theory_title": t("report.theory_title"),
            "examples": t("report.examples"),
            "recommendation": t("report.recommendation"),
            "no_issues": t("report.no_issues"),
            "expand_more": t("common.expand_more"),
            "collapse": t("common.collapse"),
            "pagespeed_screenshots": t("report.pagespeed_screenshots"),
            "homepage_screenshot_title": t("report.homepage_screenshot_title"),
            "badge_ok": t("report.badge_ok"),
            "badge_warning": t("report.badge_warning"),
            "badge_error": t("report.badge_error"),
            "badge_info": t("report.badge_info"),
            "pages_analyzed": t("report.pages_analyzed", count=audit.pages_crawled),
            "category_overview": t("report.category_overview"),
            "category_overview_desc": t("report.category_overview_description"),
            "category_label": t("report.category"),
            "status_label": t("report.status"),
            "critical_label": t("report.critical_count"),
            "warnings_label": t("report.warning_count"),
        }

        report_title_display = self._build_report_heading(t, brand)
        template_brand = {}
        if brand:
            if brand.get("logo_url"):
                template_brand["logo_url"] = brand["logo_url"]
            if brand.get("company_name"):
                template_brand["company_name"] = brand["company_name"]

        # Render template
        html = template.render(
            audit=audit,
            domain=domain,
            sections=sections,
            category_overview=category_overview,
            generated_at=datetime.now().strftime("%d.%m.%Y %H:%M"),
            SeverityLevel=SeverityLevel,
            t=translations,
            lang=lang,
            brand=template_brand,
            report_title_display=report_title_display,
            show_pages_crawled=audit.show_pages_crawled,
        )

        # Save report
        report_filename = f"audit_{audit.id}.html"
        report_path = Path(settings.REPORTS_DIR) / report_filename

        with open(report_path, "w", encoding="utf-8") as f:
            f.write(html)

        return str(report_path)

    async def generate_pdf(self, audit: AuditResult, brand: dict | None = None, show_watermark: bool = True) -> str:
        """Generate PDF report and return file path."""
        try:
            from weasyprint import HTML, CSS
        except ImportError:
            raise ImportError("weasyprint is required for PDF export. Install it with: pip install weasyprint")

        # Get translator and sections (same logic as generate())
        lang = getattr(audit, 'language', 'en') or 'en'
        t = get_translator(lang)
        domain = extract_domain(audit.url)

        sections = []
        section_order = [
            "cms", "speed", "meta_tags", "headings", "page_404",
            "images", "content", "links", "favicon", "external_links",
            "robots", "structure", "content_sections",
            "schema", "social_tags", "security", "mobile",
            "url_quality", "hreflang", "duplicates", "redirects",
        ]
        for name in section_order:
            if name in audit.results:
                result = audit.results[name]
                if lang != 'en':
                    result = translate_analyzer_content(result, lang, t)
                title = t(f"analyzers.{name}.name")
                if title == f"analyzers.{name}.name":
                    title = result.display_name
                sections.append({"id": name, "title": title, "severity": result.severity, "result": result})

        # First generate HTML (for the detailed findings)
        html_path = await self.generate(audit, brand=brand)
        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        # --- Replace summary header with cover-style content for PDF ---
        generated_at = datetime.now().strftime("%d.%m.%Y")
        report_title_display = self._build_report_heading(t, brand)
        safe_report_title = escape(report_title_display)
        cover_header = f'''
            <h1 class="pdf-cover-title">{safe_report_title}</h1>
            <div class="pdf-cover-url">Website: {domain}</div>
            <div class="pdf-cover-meta">{t("report.pages_analyzed", count=audit.pages_crawled)} · {generated_at}</div>
        '''
        # Replace the original h1 + two p tags header in the summary section
        html_content = re.sub(
            r'<h1 style="font-size: 24px; margin-bottom: 6px;">.*?</h1>\s*'
            r'<p style="color: var\(--color-text-light\); margin-bottom: 4px; font-size: 13px;">.*?</p>\s*'
            r'<p style="color: var\(--color-text-light\); margin-bottom: 24px; font-size: 13px;">.*?</p>',
            cover_header,
            html_content,
            flags=re.DOTALL,
        )

        # --- Embed branding logo as base64 so WeasyPrint doesn't need to fetch it ---
        if brand and brand.get("logo_url"):
            logo_bytes = self._fetch_logo_bytes(brand["logo_url"])
            if logo_bytes:
                import base64 as _b64
                ext = brand["logo_url"].rsplit(".", 1)[-1].lower()
                mime = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "gif": "image/gif", "webp": "image/webp"}.get(ext, "image/png")
                b64_data = _b64.b64encode(logo_bytes).decode("ascii")
                data_uri = f"data:{mime};base64,{b64_data}"
                html_content = html_content.replace(f'src="{brand["logo_url"]}"', f'src="{data_uri}"')

        # --- Change 4: Limit URL lists to 10 items for PDF ---
        html_content = self._limit_pdf_urls(html_content)

        # Create PDF
        pdf_filename = f"audit_{audit.id}.pdf"
        pdf_path = Path(settings.REPORTS_DIR) / pdf_filename

        watermark_css = """
                @bottom-right {
                    content: "seo.lvdev.co";
                    font-size: 8pt;
                    color: #D1D5DB;
                    font-family: Inter, sans-serif;
                }
        """ if show_watermark else ""

        # Print-specific CSS
        print_css_string = """
            @page {
                size: A4;
                margin: 1.5cm 1.5cm 2cm 1.5cm;
                @bottom-center {
                    content: counter(page);
                    font-size: 9pt;
                    color: #9CA3AF;
                    font-family: Inter, sans-serif;
                }
                __WATERMARK__
            }
            .sidebar {
                display: none !important;
            }
            .main {
                margin-left: 0 !important;
                padding: 0 !important;
                max-width: 100% !important;
            }
            body {
                background: white !important;
                font-size: 10pt !important;
                color: #333 !important;
            }

            /* === Cover Header (merged into first page) === */
            .pdf-cover-title {
                font-size: 22pt;
                font-weight: 700;
                color: #111827;
                margin: 0 0 8px 0;
            }
            .pdf-cover-url {
                font-size: 10pt;
                color: #111827;
                margin-bottom: 2px;
            }
            .pdf-cover-meta {
                font-size: 10pt;
                color: #111827;
                margin-bottom: 4px;
            }

            /* === Summary → Screenshot spacing === */
            #summary {
                margin-bottom: 24px !important;
            }

            /* === Category Overview === */
            #category-overview {
                page-break-after: always;
            }

            /* === Detailed Findings === */
            .section {
                margin-bottom: 18px !important;
            }
            .section-header {
                margin-bottom: 10px !important;
                padding-bottom: 8px !important;
                align-items: center !important;
            }
            .issue {
                margin-bottom: 8px !important;
            }
            .theory-block {
                margin-bottom: 10px !important;
            }
            /* Allow ALL content to split across pages — no gaps */
            .section, .section-header, .issue, .issue-header,
            .theory-block, .table-wrapper, .summary-grid, .screenshots-grid {
                page-break-inside: auto !important;
                break-inside: auto !important;
                page-break-before: auto !important;
                break-before: auto !important;
                page-break-after: auto !important;
                break-after: auto !important;
            }
            /* Summary cards: compact for A4 */
            .summary-grid {
                grid-template-columns: repeat(4, 1fr) !important;
                margin-top: 24px !important;
            }
            .summary-grid-3 {
                grid-template-columns: repeat(3, 1fr) !important;
            }
            .summary-card {
                padding: 10px !important;
                gap: 10px !important;
            }
            .summary-card .icon-circle {
                width: 32px !important;
                height: 32px !important;
                min-width: 32px !important;
            }
            .summary-card .icon-circle svg {
                width: 16px !important;
                height: 16px !important;
            }
            .summary-card .number {
                font-size: 20px !important;
            }
            .summary-card .label {
                font-size: 10px !important;
            }
            .screenshots-grid {
                grid-template-columns: 1fr !important;
            }
            .screenshot-card img {
                max-width: 100% !important;
            }
            /* Badge pill styling for WeasyPrint */
            .badge {
                display: inline-flex !important;
                align-items: center !important;
                gap: 4px !important;
                padding: 3px 10px !important;
                border-radius: 9999px !important;
                font-size: 12px !important;
                font-weight: 500 !important;
                vertical-align: middle !important;
                line-height: 1 !important;
            }
            .badge svg {
                display: inline-block !important;
                vertical-align: middle !important;
            }
            /* Force details open and hide interactive elements for PDF */
            details {
                display: block !important;
            }
            details summary {
                list-style: none !important;
                pointer-events: none !important;
            }
            details summary::marker,
            details summary::-webkit-details-marker {
                display: none !important;
            }
            details:not([open]) > *:not(summary) {
                display: block !important;
            }
            .expand-urls-btn {
                display: none !important;
            }
            .urls-hidden {
                display: block !important;
            }
            /* === Truncation: single-line URLs and table cells === */
            .data-table {
                table-layout: fixed !important;
            }
            .category-table {
                table-layout: auto !important;
            }
            .category-table td {
                white-space: normal !important;
                overflow: visible !important;
                text-overflow: clip !important;
                max-width: none !important;
            }
            .category-table .badge {
                display: inline-block !important;
                width: auto !important;
                white-space: nowrap !important;
                line-height: 1.1 !important;
                vertical-align: middle !important;
            }
            .category-table .badge svg {
                display: inline-block !important;
                vertical-align: text-bottom !important;
                margin-right: 4px !important;
            }
            .issue-urls li {
                white-space: nowrap !important;
                overflow: hidden !important;
                text-overflow: ellipsis !important;
                max-width: 100% !important;
                word-break: normal !important;
            }
            .data-table td {
                white-space: nowrap !important;
                overflow: hidden !important;
                text-overflow: ellipsis !important;
                max-width: 300px !important;
                word-break: normal !important;
            }
        """
        print_css_string = print_css_string.replace("__WATERMARK__", watermark_css)
        print_css = CSS(string=print_css_string)

        # Force all details elements to be open for PDF
        html_content = html_content.replace('<details class="theory-block">', '<details class="theory-block" open>')

        HTML(string=html_content).write_pdf(pdf_path, stylesheets=[print_css])

        return str(pdf_path)

    @staticmethod
    def _limit_pdf_urls(html: str, max_urls: int = 10) -> str:
        """Limit URL lists in PDF to max_urls items per issue."""
        def replace_ul(match):
            full = match.group(0)
            items = re.findall(r'<li>.*?</li>', full, re.DOTALL)
            if len(items) <= max_urls:
                return full
            kept = '\n'.join(items[:max_urls])
            remaining = len(items) - max_urls
            more_text = f'<li style="color: #6B7280; font-style: italic;">... and {remaining} more</li>'
            return re.sub(
                r'(<ul[^>]*>)(.*?)(</ul>)',
                lambda m: m.group(1) + '\n' + kept + '\n' + more_text + '\n' + m.group(3),
                full,
                count=0,
                flags=re.DOTALL,
            )

        return re.sub(
            r'<ul[^>]*>(?:\s*<li>.*?</li>\s*)+</ul>',
            replace_ul,
            html,
            flags=re.DOTALL,
        )

    # --- DOCX Helper Methods ---

    @staticmethod
    def _docx_set_cell_shading(cell, color_hex: str):
        """Set background color on a table cell."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        tc_pr.append(shading)

    @staticmethod
    def _docx_set_cell_left_border(cell, color_hex: str, width: str = "24"):
        """Add a colored left border to a cell."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn('w:tcBorders'))
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tc_pr.append(borders)
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), width)
        left.set(qn('w:space'), '0')
        left.set(qn('w:color'), color_hex)
        borders.append(left)

    @staticmethod
    def _docx_remove_cell_borders(cell):
        """Remove all borders from a cell."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        for side in ('w:top', 'w:left', 'w:bottom', 'w:right'):
            el = OxmlElement(side)
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')
            borders.append(el)
        tc_pr.append(borders)

    @staticmethod
    def _docx_set_cell_margins(cell, top=0, right=80, bottom=0, left=80):
        """Set cell internal margins (padding) in DXA units (1/20 of a point, ~80 DXA = 1.4mm)."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = OxmlElement('w:tcMar')
        for side, value in [('w:top', top), ('w:left', left), ('w:bottom', bottom), ('w:right', right)]:
            el = OxmlElement(side)
            el.set(qn('w:w'), str(value))
            el.set(qn('w:type'), 'dxa')
            tc_mar.append(el)
        tc_pr.append(tc_mar)

    @staticmethod
    def _docx_set_font(run, font_name: str = DOCX_DEFAULT_FONT, size_pt=None, bold=None, color_rgb=None):
        """Configure a run with font settings."""
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        run.font.name = font_name
        r = run._element
        r_pr = r.find(qn('w:rPr'))
        if r_pr is None:
            r_pr = r.makeelement(qn('w:rPr'), {})
            r.insert(0, r_pr)
        r_fonts = r_pr.find(qn('w:rFonts'))
        if r_fonts is None:
            r_fonts = r.makeelement(qn('w:rFonts'), {})
            r_pr.append(r_fonts)
        r_fonts.set(qn('w:ascii'), font_name)
        r_fonts.set(qn('w:hAnsi'), font_name)
        r_fonts.set(qn('w:cs'), font_name)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.font.bold = bold
        if color_rgb is not None:
            run.font.color.rgb = RGBColor(*color_rgb)

    @staticmethod
    def _docx_add_hyperlink(paragraph, url: str, text: str, font_name: str = DOCX_DEFAULT_FONT, font_size_pt: int = 9, color_rgb=None):
        """Add a clickable hyperlink to a Word paragraph."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt, RGBColor

        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Font
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)

        # Size
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size_pt * 2))  # Half-points
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(font_size_pt * 2))
        rPr.append(szCs)

        # Color
        rgb = color_rgb or (59, 130, 246)  # Default blue
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '{:02X}{:02X}{:02X}'.format(*rgb))
        rPr.append(color)

        # Underline
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)
        new_run_text = OxmlElement('w:t')
        new_run_text.set(qn('xml:space'), 'preserve')
        new_run_text.text = text
        new_run.append(new_run_text)
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    @staticmethod
    def _strip_docx_decorations(text: str) -> str:
        """Remove decorative symbol/emoji prefixes from DOCX text labels."""
        if not text:
            return text
        index = 0
        for ch in text:
            category = unicodedata.category(ch)
            if ch.isspace():
                index += 1
                continue
            if category[0] in ("L", "N"):
                break
            index += 1
        return text[index:].lstrip() or text

    @staticmethod
    def _fetch_logo_bytes(logo_url: str):
        """Fetch logo bytes for rendering; returns None on failure.

        Supports both HTTP(S) URLs and ``data:`` URIs (base64-encoded).
        """
        if not logo_url:
            return None
        # Handle base64 data URIs directly
        if logo_url.startswith("data:"):
            try:
                import base64 as _b64
                # data:[<mediatype>][;base64],<data>
                _, encoded = logo_url.split(",", 1)
                return _b64.b64decode(encoded)
            except Exception as exc:
                logger.warning(f"Failed to decode data-URI branding logo: {exc}")
                return None
        try:
            req = Request(logo_url, headers={"User-Agent": "seoapp-docx"})
            with urlopen(req, timeout=6) as response:
                content_type = (response.headers.get("Content-Type") or "").lower()
                if not content_type.startswith("image/"):
                    return None
                return response.read()
        except Exception as exc:
            logger.warning(f"Failed to fetch branding logo for DOCX: {exc}")
            return None

    def _docx_add_formatted_cell(self, paragraph, value, font_size_pt: int = 9):
        """Add formatted text to a table cell: colored ✓/✗/⚠ icons and clickable URLs."""
        import re as _re
        from docx.shared import RGBColor

        text = str(value) if value is not None else ""

        # Icon character → color mapping
        icon_colors = {
            '\u2713': (16, 185, 129),   # ✓ green
            '\u2714': (16, 185, 129),   # ✔ green
            '\u2717': (239, 68, 68),    # ✗ red
            '\u2718': (239, 68, 68),    # ✘ red
            '\u2716': (239, 68, 68),    # ✖ red
        }
        warning_chars = {'\u26a0'}  # ⚠
        warning_color = (245, 158, 11)  # amber

        # If the entire value is a URL, render as hyperlink
        stripped = text.strip()
        if _re.match(r'^https?://', stripped) and ' ' not in stripped:
            self._docx_add_hyperlink(paragraph, stripped, stripped, font_size_pt=font_size_pt)
            return

        # Split text into segments around icon characters
        # Build a regex pattern for all icon chars (including ⚠️ with variation selector)
        icon_pattern = _re.compile('([\u2713\u2714\u2716\u2717\u2718]|\u26a0\ufe0f?)')
        segments = icon_pattern.split(text)

        for segment in segments:
            if not segment:
                continue

            # Check if this segment is an icon character
            clean = segment.replace('\ufe0f', '')  # Remove variation selector
            if clean in icon_colors:
                run = paragraph.add_run(clean)
                self._docx_set_font(run, size_pt=font_size_pt, color_rgb=icon_colors[clean])
            elif clean in warning_chars:
                run = paragraph.add_run(clean)
                self._docx_set_font(run, size_pt=font_size_pt, color_rgb=warning_color)
            else:
                # Regular text — check if it contains a URL
                url_match = _re.search(r'(https?://[^\s]+)', segment)
                if url_match:
                    # Text before URL
                    before = segment[:url_match.start()]
                    if before:
                        run = paragraph.add_run(before)
                        self._docx_set_font(run, size_pt=font_size_pt)
                    # URL as hyperlink
                    self._docx_add_hyperlink(paragraph, url_match.group(1), url_match.group(1), font_size_pt=font_size_pt)
                    # Text after URL
                    after = segment[url_match.end():]
                    if after:
                        run = paragraph.add_run(after)
                        self._docx_set_font(run, size_pt=font_size_pt)
                else:
                    run = paragraph.add_run(segment)
                    self._docx_set_font(run, size_pt=font_size_pt)

    def _docx_parse_theory(self, doc, theory_html: str):
        """Parse theory HTML into Word paragraphs with formatting."""
        import re
        from docx.shared import Pt, RGBColor

        if not theory_html:
            return
        theory_html = unescape(theory_html)

        # Create a single-cell table for gray background
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        self._docx_remove_cell_borders(cell)
        self._docx_set_cell_shading(cell, 'F0F4F8')
        self._docx_set_cell_margins(cell, top=80, right=120, bottom=100, left=120)

        # Clear default paragraph
        cell.text = ''

        lines = theory_html.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            p = cell.add_paragraph() if cell.paragraphs[0].text or i > 0 else cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # Handle bullet points
            if line.startswith('•'):
                line = line[1:].strip()
                # Add bullet character
                run = p.add_run('  •  ')
                self._docx_set_font(run, size_pt=9, color_rgb=(107, 114, 128))

            # Parse inline HTML tags
            parts = re.split(r'(<strong>.*?</strong>|<code>.*?</code>)', line)
            for part in parts:
                if not part:
                    continue
                strong_match = re.match(r'<strong>(.*?)</strong>', part)
                code_match = re.match(r'<code>(.*?)</code>', part)
                if strong_match:
                    run = p.add_run(strong_match.group(1))
                    self._docx_set_font(run, size_pt=9, bold=True)
                elif code_match:
                    run = p.add_run(code_match.group(1))
                    self._docx_set_font(run, font_name=DOCX_MONO_FONT, size_pt=9, color_rgb=(107, 114, 128))
                else:
                    # Strip any remaining HTML
                    clean = re.sub(r'<[^>]+>', '', part)
                    if clean:
                        run = p.add_run(clean)
                        self._docx_set_font(run, size_pt=9)

    def _docx_add_issue_card(self, doc, issue, t_labels: dict):
        """Add a colored issue card as a borderless single-cell table."""
        from docx.shared import Pt, RGBColor

        severity_colors = {
            SeverityLevel.ERROR: 'FEE2E2',
            SeverityLevel.WARNING: 'FEF3C7',
            SeverityLevel.SUCCESS: 'D1FAE5',
            SeverityLevel.INFO: 'DBEAFE',
        }
        severity_text_colors = {
            SeverityLevel.ERROR: (239, 68, 68),
            SeverityLevel.WARNING: (180, 120, 0),
            SeverityLevel.SUCCESS: (16, 150, 100),
            SeverityLevel.INFO: (59, 130, 246),
        }

        bg_color = severity_colors.get(issue.severity, 'F3F4F6')
        text_color = severity_text_colors.get(issue.severity, (31, 41, 55))

        # Create single-cell table for the card
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.rows[0].cells[0]
        self._docx_remove_cell_borders(cell)
        self._docx_set_cell_shading(cell, bg_color)
        self._docx_set_cell_margins(cell, top=120, right=140, bottom=140, left=140)

        # Issue message (bold, colored)
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(issue.message)
        self._docx_set_font(run, size_pt=10, bold=True, color_rgb=text_color)

        # Details
        if issue.details:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(issue.details)
            self._docx_set_font(run, size_pt=9, color_rgb=(55, 65, 81))

        # Recommendation
        if issue.recommendation:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            rec_label = t_labels.get("recommendation", "Рекомендація")
            run = p.add_run(f"{rec_label}: ")
            self._docx_set_font(run, size_pt=9, bold=True, color_rgb=(55, 65, 81))
            run = p.add_run(issue.recommendation)
            self._docx_set_font(run, size_pt=9, color_rgb=(55, 65, 81))

        # Affected URLs
        if issue.affected_urls:
            examples_label = t_labels.get("examples", "Приклади")
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(f"{examples_label}:")
            self._docx_set_font(run, size_pt=8, bold=True, color_rgb=(75, 85, 99))
            for url in issue.affected_urls:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run("  \u2022 ")
                self._docx_set_font(run, size_pt=8, color_rgb=(55, 65, 81))
                self._docx_add_hyperlink(p, url, url, font_size_pt=8, color_rgb=(55, 65, 81))

    async def generate_docx(self, audit: AuditResult, brand: dict | None = None, show_watermark: bool = True) -> str:
        """Generate styled DOCX report and return file path."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError("python-docx is required for Word export. Install it with: pip install python-docx")

        # Setup i18n
        lang = getattr(audit, 'language', 'en') or 'en'
        t = get_translator(lang)

        t_labels = {
            "express_title": t("report.express_title"),
            "generated_at": t("report.generated_at"),
            "overview": t("report.overview"),
            "pages_crawled": t("report.pages_crawled"),
            "passed_checks": t("report.passed_checks"),
            "warnings": t("report.warnings"),
            "critical_issues": t("report.critical_issues"),
            "theory_title": t("report.theory_title"),
            "examples": t("report.examples"),
            "recommendation": t("report.recommendation"),
            "no_issues": t("report.no_issues"),
            "pagespeed_screenshots": t("report.pagespeed_screenshots"),
            "homepage_screenshot_title": t("report.homepage_screenshot_title"),
        }

        # Extract domain
        domain = extract_domain(audit.url)

        # Create document
        doc = Document()

        # --- Setup cross-platform DOCX fonts ---
        from docx.oxml import OxmlElement

        def _set_style_font(s, font_name=DOCX_DEFAULT_FONT):
            s.font.name = font_name
            rPr = s.element.get_or_add_rPr()
            r_fonts = rPr.find(qn('w:rFonts'))
            if r_fonts is None:
                r_fonts = OxmlElement('w:rFonts')
                rPr.append(r_fonts)
            r_fonts.set(qn('w:ascii'), font_name)
            r_fonts.set(qn('w:hAnsi'), font_name)
            r_fonts.set(qn('w:cs'), font_name)

        style = doc.styles['Normal']
        style.font.size = Pt(10)
        _set_style_font(style)

        for heading_level in range(1, 4):
            style_name = f'Heading {heading_level}'
            if style_name in doc.styles:
                h_style = doc.styles[style_name]
                h_style.font.color.rgb = RGBColor(31, 41, 55)
                _set_style_font(h_style)

        # Keep a fixed accent color for DOCX summaries.
        brand_primary_hex = '3B82F6'
        brand_logo_url = (brand or {}).get("logo_url")

        # =============================================
        # HEADER
        # =============================================
        if brand_logo_url:
            logo_bytes = self._fetch_logo_bytes(brand_logo_url)
            if logo_bytes:
                try:
                    from io import BytesIO
                    doc.add_picture(BytesIO(logo_bytes), width=Inches(1.5))
                except Exception as exc:
                    logger.warning(f"Failed to embed branding logo in DOCX: {exc}")

        # Title
        report_title_display = self._build_report_heading(t, brand)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.paragraph_format.space_after = Pt(2)
        run = title_para.add_run(self._strip_docx_decorations(report_title_display))
        self._docx_set_font(run, size_pt=22, bold=True, color_rgb=(17, 24, 39))

        # Website
        site_para = doc.add_paragraph()
        site_para.paragraph_format.space_before = Pt(2)
        site_para.paragraph_format.space_after = Pt(1)
        run = site_para.add_run(f"Website: {domain}")
        self._docx_set_font(run, size_pt=10, color_rgb=(17, 24, 39))

        # Meta line (pages · date)
        generated_at = datetime.now().strftime('%d.%m.%Y')
        meta_para = doc.add_paragraph()
        meta_para.paragraph_format.space_before = Pt(2)
        meta_para.paragraph_format.space_after = Pt(10)
        pages_text = t("report.pages_analyzed", count=audit.pages_crawled)
        run = meta_para.add_run(f"{pages_text} · {generated_at}")
        self._docx_set_font(run, size_pt=10, color_rgb=(107, 114, 128))

        # Summary stats table
        summary_items = [
            (t_labels['passed_checks'], str(audit.passed_checks), '10B981'),
            (t_labels['warnings'], str(audit.warnings), 'F59E0B'),
            (t_labels['critical_issues'], str(audit.critical_issues), 'EF4444'),
        ]
        if audit.show_pages_crawled:
            summary_items.insert(0, (t_labels['pages_crawled'], str(audit.pages_crawled), brand_primary_hex))

        summary_table = doc.add_table(rows=2, cols=len(summary_items))
        summary_table.style = 'Table Grid'

        for i, (label, value, color) in enumerate(summary_items):
            header_cell = summary_table.rows[0].cells[i]
            header_cell.text = ''
            p = header_cell.paragraphs[0]
            run = p.add_run(label)
            self._docx_set_font(run, size_pt=9, bold=True, color_rgb=(107, 114, 128))
            self._docx_set_cell_left_border(header_cell, color, '24')
            self._docx_set_cell_margins(header_cell, top=40, right=80, bottom=0, left=80)

            value_cell = summary_table.rows[1].cells[i]
            value_cell.text = ''
            p = value_cell.paragraphs[0]
            run = p.add_run(value)
            self._docx_set_font(run, size_pt=18, bold=True, color_rgb=(31, 41, 55))
            self._docx_set_cell_left_border(value_cell, color, '24')
            self._docx_set_cell_margins(value_cell, top=0, right=80, bottom=40, left=80)

        doc.add_paragraph()

        # Build sections list for category overview and detailed findings
        section_order = [
            "cms", "speed", "meta_tags", "headings", "page_404",
            "images", "content", "links", "favicon", "external_links",
            "robots", "structure", "content_sections",
            "schema", "social_tags", "security", "mobile",
            "url_quality", "hreflang", "duplicates", "redirects",
        ]
        docx_sections = []
        for name in section_order:
            if name in audit.results:
                result = audit.results[name]
                if lang != 'en':
                    result = translate_analyzer_content(result, lang, t)
                title = t(f"analyzers.{name}.name")
                if title == f"analyzers.{name}.name":
                    title = result.display_name
                title = self._strip_docx_decorations(title)
                docx_sections.append({"id": name, "title": title, "severity": result.severity, "result": result})

        # =============================================
        # HOMEPAGE SCREENSHOT
        # =============================================
        if audit.homepage_screenshot:
            import base64 as b64
            from io import BytesIO
            try:
                img_bytes = b64.b64decode(audit.homepage_screenshot)
                img_stream = BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(6.0))
            except Exception as e:
                logger.warning(f"Failed to add homepage screenshot to DOCX: {e}")
            doc.add_paragraph()

        # =============================================
        # CATEGORY OVERVIEW
        # =============================================
        cat_heading = doc.add_heading(t("report.category_overview"), level=1)
        cat_heading.paragraph_format.space_after = Pt(6)

        cat_desc = doc.add_paragraph()
        cat_desc.paragraph_format.space_after = Pt(10)
        run = cat_desc.add_run(t("report.category_overview_description"))
        self._docx_set_font(run, size_pt=9, color_rgb=(75, 85, 99))

        severity_order_map = {SeverityLevel.ERROR: 0, SeverityLevel.WARNING: 1, SeverityLevel.INFO: 2, SeverityLevel.SUCCESS: 3}
        sorted_sections = sorted(docx_sections, key=lambda s: severity_order_map.get(s["severity"], 4))

        cat_table = doc.add_table(rows=1 + len(sorted_sections), cols=4)
        cat_table.style = 'Table Grid'

        # Set column widths: category ~50%, others share remaining space
        cat_table.columns[0].width = Inches(3.25)
        cat_table.columns[1].width = Inches(1.25)
        cat_table.columns[2].width = Inches(1.0)
        cat_table.columns[3].width = Inches(1.0)

        # Header row
        headers = [t("report.category"), t("report.status"), t("report.critical_count"), t("report.warning_count")]
        for i, header_text in enumerate(headers):
            cell = cat_table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            if i in (1, 2, 3):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header_text)
            self._docx_set_font(run, size_pt=8, bold=True, color_rgb=(55, 65, 81))
            self._docx_set_cell_shading(cell, 'F1F5F9')

        badge_text_map = {
            SeverityLevel.SUCCESS: t("report.badge_ok"),
            SeverityLevel.WARNING: t("report.badge_warning"),
            SeverityLevel.ERROR: t("report.badge_error"),
            SeverityLevel.INFO: t("report.badge_info"),
        }
        badge_color_map = {
            SeverityLevel.SUCCESS: (6, 95, 70),
            SeverityLevel.WARNING: (146, 64, 14),
            SeverityLevel.ERROR: (153, 27, 27),
            SeverityLevel.INFO: (30, 64, 175),
        }

        for row_idx, section in enumerate(sorted_sections, 1):
            result = section["result"]
            criticals = sum(1 for iss in result.issues if iss.severity == SeverityLevel.ERROR)
            warns = sum(1 for iss in result.issues if iss.severity == SeverityLevel.WARNING)

            row = cat_table.rows[row_idx]

            if row_idx % 2 == 0:
                for shaded_cell in row.cells:
                    self._docx_set_cell_shading(shaded_cell, 'F8FAFC')

            # Category
            cell = row.cells[0]
            cell.text = ''
            run = cell.paragraphs[0].add_run(section["title"])
            self._docx_set_font(run, size_pt=8, color_rgb=(55, 65, 81))
            # Status badge
            cell = row.cells[1]
            cell.text = ''
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            badge = badge_text_map.get(section["severity"], "—")
            badge_clr = badge_color_map.get(section["severity"], (55, 65, 81))
            run = cell.paragraphs[0].add_run(badge)
            self._docx_set_font(run, size_pt=8, bold=True, color_rgb=badge_clr)
            # Critical count
            cell = row.cells[2]
            cell.text = ''
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(str(criticals))
            self._docx_set_font(run, size_pt=8, bold=criticals > 0, color_rgb=(239, 68, 68) if criticals > 0 else (156, 163, 175))
            # Warning count
            cell = row.cells[3]
            cell.text = ''
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(str(warns))
            self._docx_set_font(run, size_pt=8, bold=warns > 0, color_rgb=(245, 158, 11) if warns > 0 else (156, 163, 175))

        doc.add_paragraph()

        # =============================================
        # DETAILED FINDINGS
        # =============================================

        # --- Results Sections ---
        section_order = [
            "cms", "speed", "meta_tags", "headings", "page_404",
            "images", "content", "links", "favicon", "external_links",
            "robots", "structure", "content_sections",
            "schema", "social_tags", "security", "mobile",
            "url_quality", "hreflang", "duplicates", "redirects",
        ]

        severity_badge_text = {
            SeverityLevel.SUCCESS: (t("report.badge_ok"), (16, 185, 129)),
            SeverityLevel.WARNING: (t("report.badge_warning"), (245, 158, 11)),
            SeverityLevel.ERROR: (t("report.badge_error"), (239, 68, 68)),
            SeverityLevel.INFO: (t("report.badge_info"), (59, 130, 246)),
        }
        section_index = 0

        for name in section_order:
            if name not in audit.results:
                continue

            section_index += 1
            result = audit.results[name]

            # Translate content if needed
            if lang != 'en':
                result = translate_analyzer_content(result, lang, t)

            # Get translated section title
            section_title = t(f"analyzers.{name}.name")
            if section_title == f"analyzers.{name}.name":
                section_title = result.display_name
            section_title = self._strip_docx_decorations(section_title)

            # Section heading
            heading = doc.add_heading(f"{section_index}. {section_title}", level=1)
            heading.paragraph_format.keep_with_next = True
            heading.paragraph_format.space_before = Pt(10)
            heading.paragraph_format.space_after = Pt(4)

            # Add severity badge after heading
            badge_label, badge_color = severity_badge_text.get(
                result.severity, (t("report.badge_info"), (59, 130, 246))
            )
            run = heading.add_run(f"  ({badge_label})")
            self._docx_set_font(run, size_pt=12, bold=False, color_rgb=badge_color)

            # Summary
            if result.summary:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(result.summary)
                self._docx_set_font(run, size_pt=10, bold=True)

            # Theory section
            if result.theory:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                run = p.add_run(t_labels['theory_title'])
                self._docx_set_font(run, size_pt=10, bold=True, color_rgb=(75, 85, 99))
                self._docx_parse_theory(doc, result.theory)
                doc.add_paragraph()  # spacing after theory

            # Issues
            if result.issues:
                for issue in result.issues:
                    self._docx_add_issue_card(doc, issue, t_labels)
                    # Small spacing between cards
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(3)
                    spacer.paragraph_format.space_after = Pt(3)
            elif not result.tables:
                p = doc.add_paragraph()
                run = p.add_run(t_labels['no_issues'])
                self._docx_set_font(run, size_pt=10, color_rgb=(16, 185, 129))

            # Tables
            for table_info in result.tables:
                if not table_info.get("rows"):
                    continue

                # Table title
                table_title = table_info.get("title", "")
                if table_title:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.keep_with_next = True
                    run = p.add_run(table_title)
                    self._docx_set_font(run, size_pt=10, bold=True)

                headers = table_info.get("headers", [])
                rows = table_info.get("rows", [])

                if headers and rows:
                    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                    table.style = 'Table Grid'

                    # Header row with gray background
                    for col_idx, header in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        run = p.add_run(header)
                        self._docx_set_font(run, size_pt=8, bold=True)
                        self._docx_set_cell_shading(cell, 'F3F4F6')
                        self._docx_set_cell_margins(cell, top=50, right=80, bottom=50, left=80)

                    # Data rows
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, header in enumerate(headers):
                            value = row_data.get(header, "")
                            cell = table.rows[row_idx + 1].cells[col_idx]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            self._docx_add_formatted_cell(p, value, font_size_pt=8)
                            self._docx_set_cell_margins(cell, top=40, right=80, bottom=40, left=80)
                            # Alternating row shading
                            if row_idx % 2 == 1:
                                self._docx_set_cell_shading(cell, 'F9FAFB')

            # Embed PageSpeed screenshots if available
            if name == "speed" and result.data:
                mobile_ss = result.data.get("mobile_screenshot")
                desktop_ss = result.data.get("desktop_screenshot")
                if mobile_ss or desktop_ss:
                    import base64 as b64
                    from io import BytesIO

                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.keep_with_next = True
                    run = p.add_run(t_labels.get("pagespeed_screenshots", "PageSpeed Screenshots"))
                    self._docx_set_font(run, size_pt=10, bold=True)

                    for label, ss_data in [("Mobile", mobile_ss), ("Desktop", desktop_ss)]:
                        if ss_data:
                            p = doc.add_paragraph()
                            p.paragraph_format.space_before = Pt(4)
                            run = p.add_run(label)
                            self._docx_set_font(run, size_pt=9, bold=True, color_rgb=(75, 85, 99))

                            try:
                                img_bytes = b64.b64decode(ss_data)
                                img_stream = BytesIO(img_bytes)
                                doc.add_picture(img_stream, width=Inches(6.0))
                            except Exception as e:
                                logger.warning(f"Failed to add PageSpeed screenshot to DOCX: {e}")

            doc.add_paragraph()  # spacing between sections

        # Add footer with page numbers and branding
        from docx.oxml import OxmlElement
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Page number field
        run = footer_para.add_run()
        self._docx_set_font(run, size_pt=8, color_rgb=(156, 163, 175))
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run._element.append(fld_char_begin)

        run2 = footer_para.add_run()
        self._docx_set_font(run2, size_pt=8, color_rgb=(156, 163, 175))
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = ' PAGE '
        run2._element.append(instr_text)

        run3 = footer_para.add_run()
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run3._element.append(fld_char_end)

        if show_watermark:
            run4 = footer_para.add_run("    |    seo.lvdev.co")
            self._docx_set_font(run4, size_pt=8, color_rgb=(209, 213, 219))

        # Save document
        docx_filename = f"audit_{audit.id}.docx"
        docx_path = Path(settings.REPORTS_DIR) / docx_filename
        doc.save(docx_path)

        return str(docx_path)


def get_report_generator() -> 'ReportGenerator':
    """Get singleton ReportGenerator instance to cache Jinja2 environment."""
    global _report_generator_instance
    if _report_generator_instance is None:
        _report_generator_instance = ReportGenerator()
    return _report_generator_instance
